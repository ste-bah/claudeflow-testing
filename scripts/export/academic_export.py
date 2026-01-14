#!/usr/bin/env python3
"""
Academic Export Module

Exports assembled prose from Phase 8 to academic formats:
- LaTeX (.tex) with proper citations
- DOCX using python-docx
- Markdown with citation links

Usage:
  god export latex --input assembled.json --output paper.tex
  god export docx --input assembled.json --output paper.docx
  god export markdown --input assembled.json --output paper.md
  god export --list  # List available assembled files
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

GOD_LEARN_DIR = PROJECT_ROOT / "god-learn"
ASSEMBLED_DIR = GOD_LEARN_DIR / "assembled"


@dataclass
class Section:
    """A document section."""
    title: str
    content: str
    level: int = 1
    citations: List[str] = None

    def __post_init__(self):
        if self.citations is None:
            self.citations = []


@dataclass
class Citation:
    """A citation reference."""
    key: str
    author: str
    title: str
    year: str
    source_path: Optional[str] = None
    page: Optional[str] = None


@dataclass
class Document:
    """An assembled document."""
    title: str
    author: str
    date: str
    abstract: str
    sections: List[Section]
    citations: Dict[str, Citation]
    metadata: Dict[str, Any]


def list_assembled_files() -> List[Dict[str, Any]]:
    """List available assembled files."""
    if not ASSEMBLED_DIR.exists():
        return []

    files = []
    for f in ASSEMBLED_DIR.glob("*.json"):
        try:
            data = json.loads(f.read_text())
            files.append({
                "path": str(f),
                "filename": f.name,
                "title": data.get("title", f.stem),
                "created": data.get("created_at", "unknown"),
                "sections": len(data.get("sections", [])),
            })
        except Exception:
            continue

    files.sort(key=lambda x: x.get("created", ""), reverse=True)
    return files


def load_assembled_file(path: str) -> Optional[Document]:
    """Load an assembled JSON file."""
    file_path = Path(path)

    # Try relative to ASSEMBLED_DIR if not found
    if not file_path.exists() and not file_path.is_absolute():
        file_path = ASSEMBLED_DIR / path

    if not file_path.exists():
        return None

    try:
        data = json.loads(file_path.read_text())

        # Parse sections
        sections = []
        for s in data.get("sections", []):
            sections.append(Section(
                title=s.get("title", "Untitled"),
                content=s.get("content", ""),
                level=s.get("level", 1),
                citations=s.get("citations", [])
            ))

        # Parse citations
        citations = {}
        for key, c in data.get("citations", {}).items():
            citations[key] = Citation(
                key=key,
                author=c.get("author", "Unknown"),
                title=c.get("title", "Untitled"),
                year=c.get("year", "n.d."),
                source_path=c.get("source_path"),
                page=c.get("page")
            )

        return Document(
            title=data.get("title", "Untitled Document"),
            author=data.get("author", "Unknown Author"),
            date=data.get("date", datetime.now().strftime("%Y-%m-%d")),
            abstract=data.get("abstract", ""),
            sections=sections,
            citations=citations,
            metadata=data.get("metadata", {})
        )
    except Exception as e:
        print(f"Error loading file: {e}", file=sys.stderr)
        return None


def extract_citations_from_text(text: str, citations: Dict[str, Citation]) -> List[str]:
    """Extract citation keys from text."""
    # Match patterns like [author2023], (Smith, 2023), etc.
    patterns = [
        r'\[([a-zA-Z]+\d{4})\]',  # [author2023]
        r'\(([A-Z][a-zA-Z]+,?\s*\d{4})\)',  # (Smith, 2023)
    ]

    found = []
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for m in matches:
            # Normalize key
            key = m.replace(" ", "").replace(",", "").lower()
            if key in citations or m in citations:
                found.append(m)

    return found


# ========================
# LaTeX Export
# ========================

def escape_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    special = {
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
    }
    for char, escaped in special.items():
        text = text.replace(char, escaped)
    return text


def export_latex(doc: Document, output_path: str) -> str:
    """Export document to LaTeX format."""
    lines = []

    # Preamble
    lines.append(r"\documentclass[12pt,a4paper]{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage[T1]{fontenc}")
    lines.append(r"\usepackage{hyperref}")
    lines.append(r"\usepackage{natbib}")
    lines.append(r"\usepackage{geometry}")
    lines.append(r"\geometry{margin=1in}")
    lines.append("")

    # Title
    lines.append(r"\title{" + escape_latex(doc.title) + "}")
    lines.append(r"\author{" + escape_latex(doc.author) + "}")
    lines.append(r"\date{" + escape_latex(doc.date) + "}")
    lines.append("")
    lines.append(r"\begin{document}")
    lines.append(r"\maketitle")
    lines.append("")

    # Abstract
    if doc.abstract:
        lines.append(r"\begin{abstract}")
        lines.append(escape_latex(doc.abstract))
        lines.append(r"\end{abstract}")
        lines.append("")

    # Sections
    for section in doc.sections:
        if section.level == 1:
            lines.append(r"\section{" + escape_latex(section.title) + "}")
        elif section.level == 2:
            lines.append(r"\subsection{" + escape_latex(section.title) + "}")
        else:
            lines.append(r"\subsubsection{" + escape_latex(section.title) + "}")

        lines.append("")
        lines.append(escape_latex(section.content))
        lines.append("")

    # Bibliography
    if doc.citations:
        lines.append(r"\section*{References}")
        lines.append(r"\begin{thebibliography}{99}")
        for key, cite in sorted(doc.citations.items()):
            lines.append(r"\bibitem{" + key + "}")
            entry = f"{cite.author} ({cite.year}). \\textit{{{escape_latex(cite.title)}}}."
            if cite.page:
                entry += f" p. {cite.page}."
            lines.append(entry)
        lines.append(r"\end{thebibliography}")

    lines.append("")
    lines.append(r"\end{document}")

    # Write output
    output = "\n".join(lines)
    Path(output_path).write_text(output, encoding="utf-8")
    return output_path


# ========================
# DOCX Export
# ========================

def export_docx(doc: Document, output_path: str) -> str:
    """Export document to DOCX format."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)

    docx_doc = DocxDocument()

    # Title
    title = docx_doc.add_heading(doc.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author and date
    author_para = docx_doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"{doc.author}\n{doc.date}")
    author_run.font.size = Pt(12)

    docx_doc.add_paragraph()

    # Abstract
    if doc.abstract:
        abstract_heading = docx_doc.add_heading("Abstract", level=1)
        abstract_para = docx_doc.add_paragraph(doc.abstract)
        abstract_para.paragraph_format.first_line_indent = Inches(0.5)
        docx_doc.add_paragraph()

    # Sections
    for section in doc.sections:
        docx_doc.add_heading(section.title, level=section.level)
        if section.content:
            para = docx_doc.add_paragraph(section.content)
            para.paragraph_format.first_line_indent = Inches(0.5)

    # References
    if doc.citations:
        docx_doc.add_heading("References", level=1)
        for key, cite in sorted(doc.citations.items()):
            ref_text = f"{cite.author} ({cite.year}). {cite.title}."
            if cite.page:
                ref_text += f" p. {cite.page}."
            para = docx_doc.add_paragraph(ref_text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)

    # Save
    docx_doc.save(output_path)
    return output_path


# ========================
# Markdown Export
# ========================

def export_markdown(doc: Document, output_path: str) -> str:
    """Export document to Markdown format."""
    lines = []

    # Title
    lines.append(f"# {doc.title}")
    lines.append("")
    lines.append(f"**Author:** {doc.author}")
    lines.append(f"**Date:** {doc.date}")
    lines.append("")

    # Abstract
    if doc.abstract:
        lines.append("## Abstract")
        lines.append("")
        lines.append(doc.abstract)
        lines.append("")

    # Sections
    for section in doc.sections:
        heading = "#" * (section.level + 1)
        lines.append(f"{heading} {section.title}")
        lines.append("")
        lines.append(section.content)
        lines.append("")

    # References
    if doc.citations:
        lines.append("## References")
        lines.append("")
        for key, cite in sorted(doc.citations.items()):
            ref_text = f"- [{key}] {cite.author} ({cite.year}). *{cite.title}*."
            if cite.page:
                ref_text += f" p. {cite.page}."
            lines.append(ref_text)
        lines.append("")

    # Write output
    output = "\n".join(lines)
    Path(output_path).write_text(output, encoding="utf-8")
    return output_path


# ========================
# CLI
# ========================

def cmd_list(args) -> int:
    """List available assembled files."""
    files = list_assembled_files()

    if not files:
        print("No assembled files found.")
        print(f"Directory: {ASSEMBLED_DIR}")
        print("Run 'god assemble' to create assembled documents.")
        return 0

    print("Available Assembled Files:")
    print("-" * 70)
    print(f"{'Filename':<40} {'Sections':<10} {'Created':<20}")
    print("-" * 70)

    for f in files:
        print(f"{f['filename']:<40} {f['sections']:<10} {f['created'][:19]:<20}")

    print()
    print("Export with: god export <format> --input <filename> --output <output>")
    return 0


def cmd_export(args) -> int:
    """Export to specified format."""
    if not args.input:
        print("Error: --input is required")
        return 1

    if not args.output:
        # Generate default output name
        input_path = Path(args.input)
        ext = {"latex": ".tex", "docx": ".docx", "markdown": ".md"}.get(args.format, ".txt")
        args.output = str(input_path.stem) + ext

    doc = load_assembled_file(args.input)
    if doc is None:
        print(f"Error: Could not load file: {args.input}")
        return 1

    print(f"Exporting '{doc.title}' to {args.format}...")

    try:
        if args.format == "latex":
            output = export_latex(doc, args.output)
        elif args.format == "docx":
            output = export_docx(doc, args.output)
        elif args.format == "markdown":
            output = export_markdown(doc, args.output)
        else:
            print(f"Unknown format: {args.format}")
            return 1

        print(f"Exported to: {output}")
        return 0

    except Exception as e:
        print(f"Export failed: {e}")
        return 1


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Export assembled documents to academic formats"
    )
    ap.add_argument("format", nargs="?", choices=["latex", "docx", "markdown"],
                    help="Export format")
    ap.add_argument("--input", "-i", help="Input assembled JSON file")
    ap.add_argument("--output", "-o", help="Output file path")
    ap.add_argument("--list", action="store_true", help="List available assembled files")
    args = ap.parse_args()

    if args.list:
        return cmd_list(args)

    if not args.format:
        ap.print_help()
        print()
        print("Formats:")
        print("  latex      LaTeX document with bibliography")
        print("  docx       Microsoft Word document")
        print("  markdown   Markdown with citation links")
        print()
        print("Examples:")
        print("  god export latex --input assembled.json --output paper.tex")
        print("  god export docx -i assembled.json -o paper.docx")
        print("  god export --list")
        return 0

    return cmd_export(args)


if __name__ == "__main__":
    raise SystemExit(main())
